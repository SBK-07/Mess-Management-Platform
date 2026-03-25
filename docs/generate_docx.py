"""
Script to generate Assignment5_Git_VersionControl.docx from the markdown file.
Uses python-docx to create a properly formatted Word document.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "Assignment5_Git_VersionControl.docx")
MD_PATH     = os.path.join(os.path.dirname(__file__), "Assignment5_Git_VersionControl.md")

# ── helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def add_horizontal_line(doc):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "4F81BD")
    pb.append(bottom)
    pPr.append(pb)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)


def strip_md_inline(text):
    """Remove markdown inline syntax (bold, italic, code backticks, links)."""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*',     r'\1', text)
    text = re.sub(r'`(.+?)`',       r'\1', text)
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    return text


def add_inline_run(para, text):
    """
    Split text on **bold**, *italic*, `code` markers and add formatted runs.
    """
    pattern = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)')
    parts   = pattern.split(text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run      = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run        = para.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run            = para.add_run(part[1:-1])
            run.font.name  = 'Courier New'
            run.font.size  = Pt(9)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        else:
            para.add_run(part)


def apply_heading_style(para, level):
    colors = {
        1: RGBColor(0x1F, 0x49, 0x7D),
        2: RGBColor(0x2E, 0x74, 0xB5),
        3: RGBColor(0x5B, 0x9B, 0xD5),
    }
    for run in para.runs:
        run.font.color.rgb = colors.get(level, RGBColor(0, 0, 0))


# ── table renderer ────────────────────────────────────────────────────────────

def render_md_table(doc, lines):
    rows = []
    for line in lines:
        if re.match(r'^\s*\|[-| :]+\|\s*$', line):
            continue
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        rows.append(cells)

    if not rows:
        return

    col_count = max(len(r) for r in rows)
    # normalise row widths
    rows = [r + [''] * (col_count - len(r)) for r in rows]

    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style          = 'Table Grid'
    table.alignment      = WD_TABLE_ALIGNMENT.CENTER

    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            cell    = table.cell(ri, ci)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para    = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after  = Pt(2)

            is_header = (ri == 0)
            if is_header:
                set_cell_bg(cell, "2E74B5")
                run            = para.add_run(strip_md_inline(cell_text))
                run.bold       = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size  = Pt(9)
            else:
                add_inline_run(para, cell_text)
                for run in para.runs:
                    run.font.size = Pt(9)
                if ri % 2 == 0:
                    set_cell_bg(cell, "DCE6F1")

    doc.add_paragraph()


# ── code-block renderer ───────────────────────────────────────────────────────

def render_code_block(doc, code_lines):
    # Outer container paragraph (gives a shaded box feel via border)
    for line in code_lines:
        p = doc.add_paragraph(style='No Spacing')
        p.paragraph_format.left_indent  = Cm(0.5)
        p.paragraph_format.right_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)

        # Light-blue shading on every line
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "EBF3FB")
        pPr.append(shd)

        run            = p.add_run(line)
        run.font.name  = 'Courier New'
        run.font.size  = Pt(8.5)
        run.font.color.rgb = RGBColor(0x24, 0x29, 0x2E)

    doc.add_paragraph()


# ── main parser / builder ─────────────────────────────────────────────────────

def build_docx(md_text):
    doc = Document()

    # ── page margins ──
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.8)
        section.right_margin  = Cm(2.8)

    # ── default body font ──
    style               = doc.styles['Normal']
    style.font.name     = 'Calibri'
    style.font.size     = Pt(11)

    lines   = md_text.splitlines()
    i       = 0
    n       = len(lines)

    while i < n:
        line = lines[i]

        # ── skip the outer markdown code fence of the whole file (```markdown) ──
        if line.strip() in ('```markdown', '````markdown'):
            i += 1
            continue
        if line.strip() in ('```', '````') and i == n - 1:
            i += 1
            continue

        # ── blank line ──
        if line.strip() == '':
            i += 1
            continue

        # ── horizontal rule ──
        if re.match(r'^-{3,}$', line.strip()):
            add_horizontal_line(doc)
            i += 1
            continue

        # ── H1 title ──
        if line.startswith('# ') and not line.startswith('## '):
            text = line[2:].strip()
            p    = doc.add_heading(strip_md_inline(text), level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
                run.font.size      = Pt(20)
            i += 1
            continue

        # ── H2 ──
        if line.startswith('## '):
            text = line[3:].strip()
            p    = doc.add_heading(strip_md_inline(text), level=1)
            apply_heading_style(p, 1)
            for run in p.runs:
                run.font.size = Pt(14)
            i += 1
            continue

        # ── H3 ──
        if line.startswith('### '):
            text = line[4:].strip()
            p    = doc.add_heading(strip_md_inline(text), level=2)
            apply_heading_style(p, 2)
            for run in p.runs:
                run.font.size = Pt(12)
            i += 1
            continue

        # ── H4 ──
        if line.startswith('#### '):
            text = line[5:].strip()
            p    = doc.add_heading(strip_md_inline(text), level=3)
            apply_heading_style(p, 3)
            for run in p.runs:
                run.font.size = Pt(11)
            i += 1
            continue

        # ── Table (starts with |) ──
        if line.startswith('|'):
            table_lines = []
            while i < n and lines[i].startswith('|'):
                table_lines.append(lines[i])
                i += 1
            render_md_table(doc, table_lines)
            continue

        # ── Code block (``` or ```) ──
        if re.match(r'^```', line.strip()) or re.match(r'^````', line.strip()):
            open_fence = line.strip()
            i += 1
            code_lines = []
            while i < n:
                cur = lines[i]
                if cur.strip() in ('```', '````') or cur.strip() == open_fence:
                    i += 1
                    break
                code_lines.append(cur)
                i += 1
            render_code_block(doc, code_lines)
            continue

        # ── Bullet list ──
        if re.match(r'^[-*] ', line):
            p    = doc.add_paragraph(style='List Bullet')
            text = line[2:].strip()
            add_inline_run(p, text)
            for run in p.runs:
                if run.font.size is None:
                    run.font.size = Pt(11)
            i += 1
            continue

        # ── Numbered list ──
        if re.match(r'^\d+\. ', line):
            p    = doc.add_paragraph(style='List Number')
            text = re.sub(r'^\d+\. ', '', line).strip()
            add_inline_run(p, text)
            for run in p.runs:
                if run.font.size is None:
                    run.font.size = Pt(11)
            i += 1
            continue

        # ── Bold-only line (acts as sub-heading) ──
        if re.match(r'^\*\*.+\*\*$', line.strip()):
            p    = doc.add_paragraph()
            run  = p.add_run(strip_md_inline(line.strip()))
            run.bold      = True
            run.font.size = Pt(11)
            i += 1
            continue

        # ── Regular paragraph ──
        # Accumulate lines until a blank / special line
        para_lines = []
        while i < n:
            cur = lines[i]
            if (cur.strip() == '' or
                cur.startswith('#') or
                cur.startswith('|') or
                re.match(r'^```', cur.strip()) or
                re.match(r'^````', cur.strip()) or
                re.match(r'^-{3,}$', cur.strip()) or
                re.match(r'^[-*] ', cur) or
                re.match(r'^\d+\. ', cur)):
                break
            para_lines.append(cur.rstrip())
            i += 1

        if para_lines:
            combined = ' '.join(l for l in para_lines if l)
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            add_inline_run(p, combined)

    return doc


# ── entry point ───────────────────────────────────────────────────────────────

if __name__ == '__main__':
    with open(MD_PATH, encoding='utf-8') as f:
        md_text = f.read()

    doc = build_docx(md_text)
    doc.save(OUTPUT_PATH)
    print(f"✓ Saved: {OUTPUT_PATH}")
